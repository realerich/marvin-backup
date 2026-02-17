#!/usr/bin/env python3
"""
文档处理工具
支持PDF、Word、Excel等格式的创建、解析和转换
"""

import json
import os
from datetime import datetime
from pathlib import Path

OUTPUT_DIR = Path("/root/.openclaw/workspace/output/documents")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

class PDFProcessor:
    """PDF处理工具"""
    
    @staticmethod
    def create_simple_pdf(title, content, output_file=None):
        """创建简单PDF文档"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            
            if not output_file:
                output_file = OUTPUT_DIR / f"doc_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            
            # 尝试注册中文字体
            try:
                pdfmetrics.registerFont(TTFont('SimSun', '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc'))
                font_name = 'SimSun'
            except:
                font_name = 'Helvetica'
            
            c = canvas.Canvas(str(output_file), pagesize=letter)
            width, height = letter
            
            # 标题
            c.setFont(font_name, 20)
            c.drawString(50, height - 50, title)
            
            # 内容
            c.setFont(font_name, 12)
            y = height - 100
            for line in content.split('\n'):
                if y < 50:  # 新页面
                    c.showPage()
                    c.setFont(font_name, 12)
                    y = height - 50
                c.drawString(50, y, line[:80])  # 每行最多80字符
                y -= 20
            
            c.save()
            return str(output_file)
        
        except Exception as e:
            return f"Error: {str(e)}"
    
    @staticmethod
    def merge_pdfs(pdf_list, output_file=None):
        """合并多个PDF"""
        try:
            from PyPDF2 import PdfMerger
            
            if not output_file:
                output_file = OUTPUT_DIR / f"merged_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            
            merger = PdfMerger()
            for pdf in pdf_list:
                merger.append(pdf)
            
            merger.write(str(output_file))
            merger.close()
            return str(output_file)
        
        except Exception as e:
            return f"Error: {str(e)}"
    
    @staticmethod
    def extract_text_from_pdf(pdf_path):
        """从PDF提取文字"""
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(pdf_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            return text
        
        except Exception as e:
            return f"Error: {str(e)}"


class ExcelProcessor:
    """Excel处理工具"""
    
    @staticmethod
    def create_from_csv(csv_file, output_file=None):
        """CSV转Excel"""
        try:
            import pandas as pd
            
            if not output_file:
                output_file = OUTPUT_DIR / f"excel_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            
            df = pd.read_csv(csv_file)
            df.to_excel(output_file, index=False, engine='openpyxl')
            return str(output_file)
        
        except Exception as e:
            return f"Error: {str(e)}"
    
    @staticmethod
    def create_restaurant_excel(csv_file, output_file=None):
        """创建餐厅数据Excel，带格式"""
        try:
            import pandas as pd
            from openpyxl import load_workbook
            from openpyxl.styles import Font, PatternFill, Alignment
            
            if not output_file:
                output_file = OUTPUT_DIR / f"restaurants_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            
            # 读取CSV
            df = pd.read_csv(csv_file)
            
            # 保存为Excel
            df.to_excel(output_file, index=False, engine='openpyxl')
            
            # 添加格式
            wb = load_workbook(output_file)
            ws = wb.active
            
            # 标题行格式
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            header_font = Font(color="FFFFFF", bold=True)
            
            for cell in ws[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # 自动调整列宽
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width
            
            # 高亮高评分
            score_col = None
            for idx, cell in enumerate(ws[1], 1):
                if cell.value == '推荐分':
                    score_col = idx
                    break
            
            if score_col:
                for row in range(2, ws.max_row + 1):
                    cell = ws.cell(row=row, column=score_col)
                    try:
                        score = float(cell.value)
                        if score >= 4.5:
                            cell.fill = PatternFill(start_color="FFD700", end_color="FFD700", fill_type="solid")
                            cell.font = Font(bold=True)
                    except:
                        pass
            
            wb.save(output_file)
            return str(output_file)
        
        except Exception as e:
            return f"Error: {str(e)}"


class WordProcessor:
    """Word文档处理"""
    
    @staticmethod
    def create_summary_doc(title, sections, output_file=None):
        """创建汇总文档"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            if not output_file:
                output_file = OUTPUT_DIR / f"summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            
            doc = Document()
            
            # 标题
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 日期
            date_para = doc.add_paragraph(f'生成日期: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # 空行
            
            # 内容章节
            for section_title, section_content in sections.items():
                doc.add_heading(section_title, level=1)
                
                if isinstance(section_content, list):
                    for item in section_content:
                        doc.add_paragraph(item, style='List Bullet')
                else:
                    doc.add_paragraph(section_content)
                
                doc.add_paragraph()  # 空行
            
            doc.save(output_file)
            return str(output_file)
        
        except Exception as e:
            return f"Error: {str(e)}"


def main():
    import sys
    
    if len(sys.argv) < 2:
        print("📄 文档处理工具")
        print("\n用法:")
        print("PDF处理:")
        print("  python3 doc_tool.py pdf create <标题> <内容文件> [输出文件]")
        print("  python3 doc_tool.py pdf merge <pdf1> <pdf2> [输出文件]")
        print("  python3 doc_tool.py pdf extract <pdf文件>")
        print("\nExcel处理:")
        print("  python3 doc_tool.py excel csv <csv文件> [输出文件]")
        print("  python3 doc_tool.py excel restaurants <csv文件> [输出文件]")
        print("\nWord处理:")
        print("  python3 doc_tool.py word create <标题> <章节JSON> [输出文件]")
        print("\n示例:")
        print("  python3 doc_tool.py pdf create '测试文档' '这是一段内容'")
        sys.exit(1)
    
    cmd_type = sys.argv[1]
    
    if cmd_type == 'pdf':
        subcmd = sys.argv[2]
        
        if subcmd == 'create':
            title = sys.argv[3]
            content = sys.argv[4]
            output = sys.argv[5] if len(sys.argv) > 5 else None
            result = PDFProcessor.create_simple_pdf(title, content, output)
            print(f"✅ PDF已创建: {result}")
        
        elif subcmd == 'merge':
            pdfs = sys.argv[3:-1] if len(sys.argv) > 4 else sys.argv[3:]
            output = sys.argv[-1] if sys.argv[-1].endswith('.pdf') else None
            result = PDFProcessor.merge_pdfs(pdfs, output)
            print(f"✅ PDF已合并: {result}")
        
        elif subcmd == 'extract':
            pdf_file = sys.argv[3]
            text = PDFProcessor.extract_text_from_pdf(pdf_file)
            print(text[:1000])  # 只显示前1000字符
        
        else:
            print(f"未知PDF命令: {subcmd}")
    
    elif cmd_type == 'excel':
        subcmd = sys.argv[2]
        
        if subcmd == 'csv':
            csv_file = sys.argv[3]
            output = sys.argv[4] if len(sys.argv) > 4 else None
            result = ExcelProcessor.create_from_csv(csv_file, output)
            print(f"✅ Excel已创建: {result}")
        
        elif subcmd == 'restaurants':
            csv_file = sys.argv[3]
            output = sys.argv[4] if len(sys.argv) > 4 else None
            result = ExcelProcessor.create_restaurant_excel(csv_file, output)
            print(f"✅ 餐厅Excel已创建: {result}")
        
        else:
            print(f"未知Excel命令: {subcmd}")
    
    elif cmd_type == 'word':
        subcmd = sys.argv[2]
        
        if subcmd == 'create':
            title = sys.argv[3]
            sections = json.loads(sys.argv[4])
            output = sys.argv[5] if len(sys.argv) > 5 else None
            result = WordProcessor.create_summary_doc(title, sections, output)
            print(f"✅ Word文档已创建: {result}")
        
        else:
            print(f"未知Word命令: {subcmd}")
    
    else:
        print(f"未知类型: {cmd_type}")

if __name__ == '__main__':
    main()
